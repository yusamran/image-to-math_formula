import os
import requests
import streamlit as st
from PIL import Image
from io import BytesIO
from docx import Document

# ----------------------------
# ✅ Local torch cache config
# ----------------------------
os.environ["TORCH_HOME"] = "./torch_cache"

# ✅ USE YOUR HUGGING FACE LINK HERE!
weights_url = "https://huggingface.co/yusamran/latex-ocr-weights/resolve/main/weights.pth"
weights_path = "./torch_cache/hub/checkpoints/weights.pth"

# ✅ Download weights if missing
if not os.path.exists(weights_path):
    st.info("⬇️ Downloading model weights from Hugging Face (~50MB)...")
    os.makedirs(os.path.dirname(weights_path), exist_ok=True)
    with requests.get(weights_url, stream=True) as r:
        r.raise_for_status()
        with open(weights_path, "wb") as f:
            for chunk in r.iter_content(chunk_size=8192):
                f.write(chunk)
    st.success("✅ Weights downloaded!")

# ✅ Monkey-patch download_checkpoints() to prevent writing to site-packages
import pix2tex.model.checkpoints.get_latest_checkpoint as glc
glc.download_checkpoints = lambda: None

# ✅ Import LatexOCR
from pix2tex.cli import LatexOCR

# ----------------------------
# Streamlit App UI
# ----------------------------
st.title("🧮 Free Image-to-LaTeX Converter (pix2tex)")

uploaded_file = st.file_uploader(
    "Upload a formula image",
    type=["png", "jpg", "jpeg", "bmp", "gif", "webp"]
)

if uploaded_file:
    image = Image.open(uploaded_file).convert("RGB")
    st.image(image, caption="Uploaded Image", use_column_width=True)

    if st.button("Convert to LaTeX"):
        st.info("⏳ Processing image...")
        model = LatexOCR(arguments=None)
        model.args.checkpoint = weights_path  # Force local weights
        model.model.load_state_dict(
            torch.load(weights_path, map_location=model.args.device)
        )
        latex_result = model(image)

        st.success("✅ Recognized LaTeX:")
        st.code(latex_result, language="latex")

        # Export to Word
        doc = Document()
        doc.add_paragraph("Recognized LaTeX formula:")
        doc.add_paragraph(latex_result)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📄 Download Word File",
            data=buffer,
            file_name="formula.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.info("ℹ️ Allowed: png, jpg, jpeg, bmp, gif, webp")
